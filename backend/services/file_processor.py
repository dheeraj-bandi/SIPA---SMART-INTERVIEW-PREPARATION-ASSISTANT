import os
import fitz  # PyMuPDF
import docx
import logging
from typing import Optional
import tempfile
import shutil

logger = logging.getLogger(__name__)

class FileProcessor:
    def __init__(self):
        """Initialize file processor"""
        self.supported_extensions = ['.pdf', '.docx', '.txt']
        
    def extract_text(self, file_path: str) -> str:
        """Extract text from various file formats"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text: {str(e)}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
                text += "\n"  # Add page break
            
            doc.close()
            
            if not text.strip():
                raise Exception("No text content found in PDF")
            
            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error reading PDF file: {str(e)}")
            raise Exception(f"Failed to read PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise Exception("No text content found in DOCX")
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error reading DOCX file: {str(e)}")
            raise Exception(f"Failed to read DOCX: {str(e)}")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if not text.strip():
                raise Exception("No text content found in TXT file")
            
            logger.info(f"Successfully extracted {len(text)} characters from TXT")
            return text.strip()
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                logger.info(f"Successfully extracted {len(text)} characters from TXT (latin-1)")
                return text.strip()
            except Exception as e:
                logger.error(f"Error reading TXT file with latin-1 encoding: {str(e)}")
                raise Exception(f"Failed to read TXT file: {str(e)}")
        except Exception as e:
            logger.error(f"Error reading TXT file: {str(e)}")
            raise Exception(f"Failed to read TXT: {str(e)}")
    
    def validate_file_size(self, file_path: str, max_size_mb: int = 16) -> bool:
        """Validate file size"""
        try:
            file_size = os.path.getsize(file_path)
            max_size_bytes = max_size_mb * 1024 * 1024
            
            if file_size > max_size_bytes:
                logger.warning(f"File size {file_size} bytes exceeds limit of {max_size_bytes} bytes")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"Error checking file size: {str(e)}")
            return False
    
    def validate_file_type(self, file_path: str) -> bool:
        """Validate file type"""
        file_extension = os.path.splitext(file_path)[1].lower()
        return file_extension in self.supported_extensions
    
    def clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        import re
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[^\w\s\-\.\,\;\:\(\)\[\]\/\@\+\#\%\&\*]', '', text)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def create_temp_file(self, original_filename: str) -> str:
        """Create a temporary file path"""
        temp_dir = tempfile.gettempdir()
        temp_filename = f"resume_analyzer_{os.urandom(8).hex()}_{original_filename}"
        return os.path.join(temp_dir, temp_filename)
    
    def cleanup_temp_files(self, file_paths: list):
        """Clean up temporary files"""
        for file_path in file_paths:
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                    logger.info(f"Cleaned up temporary file: {file_path}")
            except Exception as e:
                logger.warning(f"Failed to clean up {file_path}: {str(e)}")
    
    def get_file_info(self, file_path: str) -> dict:
        """Get file information"""
        try:
            stat = os.stat(file_path)
            return {
                'filename': os.path.basename(file_path),
                'size_bytes': stat.st_size,
                'size_mb': round(stat.st_size / (1024 * 1024), 2),
                'extension': os.path.splitext(file_path)[1].lower(),
                'created': stat.st_ctime,
                'modified': stat.st_mtime
            }
        except Exception as e:
            logger.error(f"Error getting file info: {str(e)}")
            return {}
    
    def convert_to_text_file(self, source_path: str, output_path: str) -> bool:
        """Convert extracted text to a text file"""
        try:
            text = self.extract_text(source_path)
            cleaned_text = self.clean_extracted_text(text)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(cleaned_text)
            
            logger.info(f"Successfully converted to text file: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error converting to text file: {str(e)}")
            return False